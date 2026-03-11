"""Generate the technical readout (.docx) from structured data sources.

Output: outputs/technical_readout.docx
"""

from datetime import date

try:
    from docx import Document
    from docx.shared import Pt, RGBColor, Inches
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    from docx.enum.table import WD_TABLE_ALIGNMENT
    from docx.oxml.ns import qn
except ImportError:
    raise SystemExit("python-docx not installed. Run: uv sync")

from load_data import (
    load_trials, load_memberships, load_fields, load_criteria,
    load_not_evaluable, unique_trial_count, trials_per_condition,
    recruiting_trials, active_trials, condition_label, ensure_outputs,
    OUTPUTS,
)

NAVY = RGBColor(0x1B, 0x3A, 0x5C)
ACCENT = RGBColor(0x2E, 0x75, 0xB6)
WHITE = RGBColor(0xFF, 0xFF, 0xFF)


def _shade_cell(cell, color_hex: str):
    tcPr = cell._element.get_or_add_tcPr()
    shd = tcPr.makeelement(qn("w:shd"), {
        qn("w:val"): "clear", qn("w:color"): "auto", qn("w:fill"): color_hex,
    })
    tcPr.append(shd)


def add_styled_table(doc, headers, rows, col_widths=None):
    table = doc.add_table(rows=1 + len(rows), cols=len(headers))
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.style = "Table Grid"
    for i, h in enumerate(headers):
        cell = table.rows[0].cells[i]
        cell.text = ""
        p = cell.paragraphs[0]
        run = p.add_run(h)
        run.bold = True
        run.font.size = Pt(10)
        run.font.color.rgb = WHITE
        run.font.name = "Calibri"
        _shade_cell(cell, "1B3A5C")
    for r_idx, row in enumerate(rows):
        for c_idx, val in enumerate(row):
            cell = table.rows[r_idx + 1].cells[c_idx]
            cell.text = ""
            p = cell.paragraphs[0]
            run = p.add_run(str(val))
            run.font.size = Pt(9)
            run.font.name = "Calibri"
            if r_idx % 2 == 1:
                _shade_cell(cell, "F2F7FB")
    if col_widths:
        for row in table.rows:
            for i, w in enumerate(col_widths):
                row.cells[i].width = Inches(w)
    return table


def add_callout(doc, text, bg="E8F0F7"):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(6)
    p.paragraph_format.space_after = Pt(6)
    pPr = p._element.get_or_add_pPr()
    shd = pPr.makeelement(qn("w:shd"), {
        qn("w:val"): "clear", qn("w:color"): "auto", qn("w:fill"): bg,
    })
    pPr.append(shd)
    run = p.add_run(text)
    run.font.size = Pt(10)
    run.font.name = "Calibri"
    run.font.color.rgb = NAVY
    return p


def generate():
    trials = load_trials()
    memberships = load_memberships()
    fields = load_fields()
    criteria = load_criteria()
    not_eval = load_not_evaluable()
    ensure_outputs()

    n_unique = unique_trial_count(trials)
    n_recruiting = len(recruiting_trials(trials))
    n_active = len(active_trials(trials))
    cond_counts = trials_per_condition(memberships)

    doc = Document()
    for section in doc.sections:
        section.page_width = Inches(8.5)
        section.page_height = Inches(11)
        section.top_margin = Inches(1)
        section.bottom_margin = Inches(0.75)
        section.left_margin = Inches(1)
        section.right_margin = Inches(1)

    style = doc.styles["Normal"]
    style.font.name = "Calibri"
    style.font.size = Pt(11)
    style.paragraph_format.space_after = Pt(6)
    style.paragraph_format.line_spacing = 1.15

    for level, (size, color) in enumerate([
        (Pt(26), NAVY), (Pt(18), ACCENT), (Pt(14), ACCENT),
    ], start=1):
        h = doc.styles[f"Heading {level}"]
        h.font.name = "Calibri"
        h.font.size = size
        h.font.color.rgb = color
        h.font.bold = True

    # ── Cover ──
    doc.add_paragraph()
    doc.add_paragraph()
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run("iDHEA PRIMARY EYE CARE DATASET")
    run.font.size = Pt(28)
    run.font.color.rgb = NAVY
    run.bold = True

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run("Clinical Trial Pre-Screening Analysis")
    run.font.size = Pt(22)
    run.font.color.rgb = ACCENT

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run("Technical Readout for Data Science Team")
    run.font.size = Pt(16)
    run.font.color.rgb = RGBColor(0x66, 0x66, 0x66)

    doc.add_paragraph()
    add_callout(doc,
        f"Date: {date.today().strftime('%B %d, %Y')}  |  "
        "Classification: Internal  |  "
        "Data Source: iDHEA Primary Eye Care Data Dictionary  |  "
        "Trial Source: ClinicalTrials.gov (Phase II/III, all statuses)")
    doc.add_page_break()

    # ── Executive Summary ──
    doc.add_heading("1. Executive Summary", level=1)
    doc.add_paragraph(
        f"This document presents a systematic mapping of the iDHEA Primary Eye Care "
        f"dataset against {n_unique} unique Phase II/III clinical trials registered on "
        f"ClinicalTrials.gov. The analysis identifies which iDHEA data fields can be "
        f"used to computationally pre-screen patients for anatomical trial criteria "
        f"across {len(cond_counts)} ophthalmic condition categories."
    )
    add_callout(doc,
        f"KEY FINDING: iDHEA's existing OCT analysis fields (ETDRS grid, fluid "
        f"detection, RNFL/GCL thickness, RPE loss) can pre-screen for anatomical "
        f"criteria in {n_unique} trials. Of these, {n_recruiting} are currently "
        f"recruiting. Pre-screening does not replace clinical eligibility assessment "
        f"-- BCVA, IOP, HbA1c, and treatment history require site-level data.",
        bg="DFF0D8")

    # ── Field Catalog ──
    doc.add_heading("2. iDHEA Field Catalog", level=1)
    doc.add_heading("2.1 Fields with Direct Pre-Screening Use", level=2)
    direct_fields = [f for f in fields if f["confidence_for_prescreening"] == "direct"]
    add_styled_table(doc,
        ["Field", "Modality", "Direct Use", "Limitations"],
        [[f["field_name"], f["modality"], f["direct_use"],
          f["limitations"][:100]] for f in direct_fields],
        col_widths=[1.5, 1.0, 2.5, 1.5])

    doc.add_heading("2.2 Fields with Partial Pre-Screening Use", level=2)
    partial_fields = [f for f in fields if f["confidence_for_prescreening"] == "partial"]
    add_styled_table(doc,
        ["Field", "Modality", "Direct Use", "Limitations"],
        [[f["field_name"], f["modality"], f["direct_use"],
          f["limitations"][:100]] for f in partial_fields],
        col_widths=[1.5, 1.0, 2.5, 1.5])

    doc.add_heading("2.3 Future Augmentation (Not Currently Evaluable)", level=2)
    future_fields = [f for f in fields if f["confidence_for_prescreening"] == "not_evaluable"]
    add_styled_table(doc,
        ["Field", "Modality", "Potential Use", "Limitations"],
        [[f["field_name"], f["modality"], f["direct_use"],
          f["limitations"][:100]] for f in future_fields],
        col_widths=[1.5, 1.0, 2.5, 1.5])
    doc.add_paragraph(
        "Note: RETFound embeddings require a validated downstream classifier "
        "to be useful for pre-screening. No such classifier currently exists "
        "in the iDHEA pipeline. These are positioned as augmentation opportunities."
    )

    doc.add_heading("2.4 Data Gaps (Not Available in iDHEA)", level=2)
    add_styled_table(doc,
        ["Missing Field", "Trial Prevalence", "Impact", "Remediation Path"],
        [[ne["description"], ne["trial_prevalence"], ne["impact"][:80],
          ne["remediation"]] for ne in not_eval],
        col_widths=[1.5, 1.2, 2.0, 1.8])

    doc.add_page_break()

    # ── Criteria Mapping ──
    doc.add_heading("3. Criteria Mapping", level=1)
    for conf, label, desc in [
        ("direct", "Direct", "iDHEA field maps to criterion without external data"),
        ("partial", "Partial", "Suggestive but requires clinical confirmation"),
        ("not_evaluable", "Not Evaluable", "Requires data outside iDHEA"),
    ]:
        relevant = [c for c in criteria if c["confidence"] == conf]
        if not relevant:
            continue
        doc.add_heading(f"3.{['direct','partial','not_evaluable'].index(conf)+1} "
                        f"{label} ({desc})", level=2)
        add_styled_table(doc,
            ["Criterion", "Type", "iDHEA Fields", "External Deps"],
            [[c["criterion_text"][:60], c["criterion_type"],
              ", ".join(c["idhea_fields"]) or "None",
              ", ".join(c["external_dependencies"]) or "None"]
             for c in relevant],
            col_widths=[2.5, 0.8, 1.5, 1.7])

    doc.add_page_break()

    # ── Trial Coverage ──
    doc.add_heading("4. Trial Coverage by Condition", level=1)
    doc.add_paragraph(
        f"{n_unique} unique trials mapped across {len(cond_counts)} conditions. "
        f"A trial may appear in multiple condition categories."
    )
    coverage_rows = []
    for cond_key, count in sorted(cond_counts.items(), key=lambda x: -x[1]):
        coverage_rows.append([
            condition_label(cond_key), str(count),
            str(len([t for t in trials
                     if t["nct_id"] in {m["nct_id"] for m in memberships
                                        if m["condition_category"] == cond_key}
                     and t["status"] == "RECRUITING"])),
        ])
    add_styled_table(doc,
        ["Condition", "Trials (incl. overlaps)", "Recruiting"],
        coverage_rows,
        col_widths=[2.5, 2.0, 2.0])

    # ── Methodology ──
    doc.add_heading("5. Methodology", level=1)
    doc.add_paragraph(
        "Trials were identified via the ClinicalTrials.gov API v2 using "
        "condition-specific searches filtered to Phase II and Phase III "
        "studies. Deduplication was performed on NCT ID. Condition membership "
        "is tracked separately to prevent conflation of unique trial counts "
        "with per-condition counts."
    )

    # ── Next Steps ──
    doc.add_heading("6. Recommended Next Steps", level=1)
    steps = [
        "Quantify matchable population: run CST + fluid + demographic queries "
        "against full iDHEA dataset for top 10 recruiting trials.",
        "Build eligibility rules engine: codify criteria as computable rules "
        "against iDHEA fields, starting with CST thresholds + fluid detection.",
        "Pilot RETFound classifier: train on embeddings for DR severity grading "
        "to close the DRSS gap.",
        "Internal demo: dashboard showing active trials, iDHEA population match "
        "per trial, per-site distribution, and data gap flags.",
    ]
    for i, step in enumerate(steps, 1):
        p = doc.add_paragraph()
        run = p.add_run(f"Step {i}: ")
        run.bold = True
        run.font.color.rgb = NAVY
        p.add_run(step)

    # Footer
    doc.add_paragraph()
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run("CONFIDENTIAL -- iDHEA Internal Use Only")
    run.font.size = Pt(9)
    run.font.color.rgb = RGBColor(0x99, 0x99, 0x99)
    run.italic = True

    out_path = OUTPUTS / "technical_readout.docx"
    doc.save(str(out_path))
    print(f"Wrote {out_path}")


if __name__ == "__main__":
    generate()
